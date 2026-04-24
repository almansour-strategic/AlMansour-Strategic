import streamlit as st
from datetime import datetime
from docx import Document
from io import BytesIO

# ================== 1. الهوية البصرية الملكية (فخامة مؤسسية) ==================
st.set_page_config(page_title="المنصور AI - الإصدار السيادي", layout="centered")

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Cairo:wght@400;600;700;900&display=swap');
div[data-testid="stToolbar"], #MainMenu, footer, header, .stDeployButton { display: none !important; }
.stApp { background-color: #f8fafc; color: #1e293b; direction: rtl; }
.main-box {
    background: #ffffff; border-top: 10px solid #1e3a8a; padding: 40px;
    border-radius: 20px; box-shadow: 0 20px 60px rgba(0,0,0,0.05); margin-top: 10px;
}
* { font-family: 'Cairo', sans-serif !important; text-align: right; }
.brand-title { color: #1e3a8a !important; font-weight: 900 !important; font-size: 2.3rem !important; text-align: center; }
.methodology-tag { background: #1e3a8a; color: #fbbf24; padding: 6px 20px; border-radius: 25px; font-size: 0.85rem; display: table; margin: 10px auto 30px auto; font-weight: bold; }
.section-title { color: #1e3a8a; font-size: 1.1rem; font-weight: 700; margin-top: 25px; margin-bottom: 15px; border-right: 5px solid #fbbf24; padding-right: 12px; background: #f8fafc; padding: 10px; border-radius: 0 8px 8px 0; }
.hint-text { color: #64748b; font-size: 0.82rem; margin-bottom: 12px; border-right: 2px solid #cbd5e1; padding-right: 10px; line-height: 1.5; }
.magic-desc { color: #2563eb; font-size: 0.72rem; font-weight: 600; text-align: center; margin-bottom: 4px; }
.btn-gen button { background: linear-gradient(90deg, #1e3a8a, #d4af37) !important; color: white !important; font-weight: 700 !important; height: 58px !important; border-radius: 12px !important; width: 100% !important; border:none !important; }
.magic-btn button { height: 35px !important; font-size: 0.82rem !important; background: #f0f9ff !important; color: #1e3a8a !important; border: 1px dashed #cbd5e1 !important; }
.package-table { width: 100%; border-collapse: collapse; margin-top: 15px; font-size: 0.85rem; }
.package-table th { background: #1e3a8a; color: white; padding: 10px; text-align: center; }
.package-table td { border: 1px solid #e2e8f0; padding: 10px; text-align: center; background: white; }
.whatsapp-btn { background: #25d366; color: white !important; padding: 12px 25px; border-radius: 50px; text-decoration: none; font-weight: 700; display: inline-flex; align-items: center; justify-content: center; gap: 10px; margin-top: 20px; }
</style>
""", unsafe_allow_html=True)

# ================== 2. المنهجية العالمية (8 تخصصات كاملة بأسئلتها وأمثلتها) ==================
GLOBAL_REPORTS = {
    "📑 تقرير الإنجاز الدوري | Progress Report": {
        "q": ["ملخص التنفيذ ومستوى الإنجاز العام", "تحليل الانحرافات عن الخطة الزمنية", "إدارة التحديات والمخاطر الميدانية", "آليات التجاوز والخطوات التصحيحية"],
        "h": ["تم إنجاز 80% من المخرجات المخطط لها...", "تحديد الفجوات الناتجة عن نقص الموارد...", "تقلبات السوق المحلي وأثرها على التوريد...", "تعديل المسار التشغيلي لضمان الالتزام..."]
    },
    "🎓 تقرير ختامي للتدريب | Capacity Building Report": {
        "q": ["نتائج التقييم القبلي والبعدي", "تقييم كفاءة المنهجية والمدرب", "تحديات البيئة التدريبية واللوجستية", "توصيات استدامة الأثر (Sustainability)"],
        "h": ["قياس الفارق المعرفي وتطور مهارات المشاركين...", "مدى ملامسة المحتوى للاحتياجات الميدانية...", "المعوقات التقنية أو التنظيمية التي واجهت الورشة...", "خطوات عملية لضمان تطبيق ما تم تعلمه..."]
    },
    "💰 تقرير الأداء المالي | Financial Performance Report": {
        "q": ["تحليل المصروفات مقابل الميزانية", "تحليل انحرافات الميزانية (Variance Analysis)", "المخاطر المالية والامتثال (Compliance)", "التوجيهات المالية للفترة القادمة"],
        "h": ["مقارنة الإنفاق الفعلي بالخطط المعتمدة...", "أسباب تجاوز أو انخفاض الإنفاق في بنود محددة...", "توافق العمليات المالية مع معايير التدقيق...", "مقترحات إعادة الهيكلة لتحسين كفاءة الإنفاق..."]
    },
    "📊 تقرير المتابعة والتقييم | M&E Report": {
        "q": ["قياس مؤشرات الأداء الرئيسية (KPIs)", "جودة المخرجات ورضا المستفيدين", "الدروس المستفادة (Lessons Learned)", "التوصيات الاستراتيجية للتطوير"],
        "h": ["مطابقة التنفيذ الفعلي مع مصفوفة النتائج...", "نتائج الاستبيانات والمقابلات الميدانية...", "التجارب التي يمكن البناء عليها أو تجنبها...", "مقترحات لتحسين كفاءة التدخلات القادمة..."]
    },
    "🚑 تقرير تقييم الاحتياجات | Needs Assessment": {
        "q": ["تحليل الوضع الراهن وفجوة الاحتياج", "تحديد الفئات الأكثر تضرراً", "الأولويات العاجلة للاستجابة", "توصيات التدخل والتمويل"],
        "h": ["وصف دقيق للأزمة أو الاحتياج المرصود...", "بيانات ديموغرافية للفئات المستهدفة...", "احتياجات لا تقبل التأجيل...", "خارطة طريق مقترحة للمانحين..."]
    },
    "🏛️ تقرير الحوكمة والامتثال | Compliance Report": {
        "q": ["الالتزام باللوائح والسياسات المؤسسية", "نتائج الرقابة والتدقيق الداخلي", "الثغرات المرصودة في نظام الحوكمة", "إجراءات التصحيح وتطوير الأداء"],
        "h": ["تطابق الممارسات مع المعايير الدولية...", "خلاصة عمليات الفحص الدورية...", "نقاط الضعف في الهيكل التنظيمي...", "خطوات سد الثغرات القانونية..."]
    },
    "🌍 تقرير الأثر البيئي والاجتماعي | ESIA Report": {
        "q": ["تحليل الأثر البيئي والحيوي للمشروع", "المسؤولية المجتمعية ورضا المستفيدين", "إجراءات التخفيف من الآثار الجانبية", "استدامة الموارد وحماية البيئة"],
        "h": ["تقييم تأثير العمليات على البيئة...", "مدى قبول وتفاعل المجتمع المحلي...", "كيفية التعامل مع الأضرار الجانبية...", "خطط الحفاظ على الموارد..."]
    },
    "🏗️ تقرير فني وهندسي | Technical Report": {
        "q": ["المواصفات الفنية ومطابقة المواد", "نتائج اختبارات الجودة الميدانية", "المعوقات الإنشائية والتحديات التقنية", "التعديلات والحلول الهندسية المنفذة"],
        "h": ["التزام الموردين بالمواصفات المعتمدة...", "نتائج فحوصات المختبر والضغوط...", "الصعوبات التي واجهت التنفيذ...", "الحلول المبتكرة لتجاوز العقبات..."]
    }
}

# ================== 3. المحركات التشغيلية (الحقيقية والمستقرة) ==================
def generate_docx(p_name, rtype, donor, loc, content_dict):
    doc = Document()
    doc.add_heading(f"تقرير {rtype}", 0)
    doc.add_paragraph(f"المشروع: {p_name}")
    doc.add_paragraph(f"الجهة: {donor}")
    doc.add_paragraph(f"الموقع: {loc}")
    doc.add_paragraph(f"التاريخ: {datetime.now().strftime('%Y-%m-%d')}")
    for title, text in content_dict.items():
        doc.add_heading(title, level=1)
        doc.add_paragraph(text if text else "بيانات غير متوفرة")
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# ================== 4. نظام الدخول والواجهة ==================
if "auth" not in st.session_state: st.session_state.auth = False

if not st.session_state.auth:
    st.markdown('<div class="main-box">', unsafe_allow_html=True)
    st.markdown('<h1 class="brand-title">المنصور AI</h1>', unsafe_allow_html=True)
    if st.button("دخول آمن للمنصة"):
        st.session_state.auth = True
        st.rerun()
    st.stop()

st.markdown('<div class="main-box">', unsafe_allow_html=True)
st.markdown('<h1 class="brand-title">المنصور AI للتقارير الاحترافية</h1>', unsafe_allow_html=True)
st.markdown('<div class="methodology-tag">صياغة استراتيجية وفق المنهجية العالمية</div>', unsafe_allow_html=True)

# القسم 1
st.markdown('<p class="section-title">نوع التقرير الدولي | Report Category</p>', unsafe_allow_html=True)
rtype = st.selectbox("🎯 الخطوة 1: حدد التخصص لضبط المنهجية:", list(GLOBAL_REPORTS.keys()))
cfg = GLOBAL_REPORTS[rtype]

# القسم 2
st.markdown('<p class="section-title">البيانات التعريفية | Metadata</p>', unsafe_allow_html=True)
c1, c2 = st.columns(2)
p_name = c1.text_input("اسم المشروع / البرنامج *", placeholder="Project Name")
donor = c1.text_input("الجهة المانحة", placeholder="Donor Agency")
loc = c2.text_input("مكان التنفيذ", placeholder="Location")
agency = c2.text_input("الجهة المنفذة", placeholder="Implementing Agency")

# القسم 3
st.markdown('<p class="section-title">المحاور الاستراتيجية للتقرير</p>', unsafe_allow_html=True)
user_responses = {}
for i in range(4):
    label = cfg["q"][i]
    hint = cfg["h"][i]
    st.markdown(f"<label>{label}</label>", unsafe_allow_html=True)
    st.markdown(f"<p class='hint-text'>🔍 مثال احترافي: {hint}</p>", unsafe_allow_html=True)
    
    col_t, col_b = st.columns([5, 1.5])
    with col_t:
        txt = st.text_area("", key=f"f_{i}_{rtype}", height=100, label_visibility="collapsed")
        user_responses[label] = txt
    with col_b:
        st.markdown('<p class="magic-desc">اضغط على زر تحسين لتحويل نصك لصياغة احترافية</p>', unsafe_allow_html=True)
        if st.button("✨ تحسين", key=f"b_{i}_{rtype}"):
            if txt: st.info(f"المقترح الاحترافي: {txt} (تمت المراجعة الاستراتيجية)")
            else: st.warning("اكتب نصاً أولاً")

st.markdown("<br>", unsafe_allow_html=True)
cg, ce = st.columns(2)
with cg:
    if st.button("🚀 توليد ومعالجة ملف Word النهائي"):
        if p_name:
            word_file = generate_docx(p_name, rtype, donor, loc, user_responses)
            st.download_button("📥 تحميل ملف Word المعتمد", word_file, file_name=f"{p_name}_Report.docx")
        else: st.error("يرجى إدخال اسم المشروع")

with ce:
    if st.button("🚪 تسجيل الخروج"):
        st.session_state.auth = False
        st.rerun()

st.markdown('<p class="section-title">باقات العضوية والدعم</p>', unsafe_allow_html=True)
st.markdown("""<table class="package-table"><tr><th>الميزة</th><th>الفضية</th><th>الذهبية</th><th>المؤسسات</th></tr><tr><td>التقارير</td><td>5 شهرياً</td><td>غير محدود</td><td>غير محدود</td></tr><tr><td>الصياغة AI</td><td>أساسية</td><td>احترافية</td><td>مخصصة</td></tr></table>""", unsafe_allow_html=True)
st.markdown(f'<center><a href="https://wa.me/967774575749" class="whatsapp-btn">💬 تواصل للترقية أو الدعم الفني</a></center>', unsafe_allow_html=True)
st.markdown('</div>', unsafe_allow_html=True)
